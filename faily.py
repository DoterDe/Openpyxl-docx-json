from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Border, Side



wb1 = Workbook()
ws1 = wb1.active

ws1.cell(row=1, column=1, value='11111')
wb1.save('file1.xlsx')
wb2 = Workbook()
ws2 = wb2.active

ws2.cell(row=1, column=1, value='22222')
wb2.save('file2.xlsx')
wb3 = Workbook()
ws3 = wb3.active

ws3.cell(row=1, column=1, value='33333')
wb3.save('file3.xlsx')


wb1 = load_workbook('file1.xlsx')
wb2 = load_workbook('file2.xlsx')
wb3 = load_workbook('file3.xlsx')

data = []
for wb in [wb1, wb2, wb3]:
    ws = wb.active
    for row in ws.iter_rows(values_only=True):
        data.extend(row)


sorted_data = sorted(data, reverse=True)


wb = Workbook()
ws = wb.active


font = Font(name='Calibri', bold=True)
border = Border(left=Side(style='thin'), 
                right=Side(style='thin'), 
                top=Side(style='thin'), 
                bottom=Side(style='thin'))

for i, value in enumerate(sorted_data):
    cell = ws.cell(row=i+1, column=1, value=value)
    cell.font = font
    cell.border = border

wb.save('output.xlsx')



import json
date=[
    {
        "userId": 1,
        "id": 1,
        "title": "delectus aut autem",
        "completed": False
    },
    {
        "userId": 1,
        "id": 2,
        "title": "quis ut nam facilis et officia qui",
        "completed": False
    },
    {
        "userId": 1,
        "id": 3,
        "title": "fugiat veniam minus",
        "completed": False
    },
    {
        "userId": 1,
        "id": 4,
        "title": "et porro tempora",
        "completed": True
    },
    {
        "userId": 1,
        "id": 5,
        "title": "laboriosam mollitia et enim quasi adipisci quia provident illum",
        "completed": False
    },
    {
        "userId": 1,
        "id": 6,
        "title": "qui ullam ratione quibusdam voluptatem quia omnis",
        "completed": False
    },
    {
        "userId": 1,
        "id": 7,
        "title": "illo expedita consequatur quia in",
        "completed": False
    },
    {
        "userId": 1,
        "id": 8,
        "title": "quo adipisci enim quam ut ab",
        "completed": True
    },
    {
        "userId": 1,
        "id": 9,
        "title": "molestiae perspiciatis ipsa",
        "completed": False
    },
    {
        "userId": 1,
        "id": 10,
        "title": "illo est ratione doloremque quia maiores aut",
        "completed": True
    },
    {
        "userId": 1,
        "id": 11,
        "title": "vero rerum temporibus dolor",
        "completed": True
    },
    {
        "userId": 1,
        "id": 12,
        "title": "ipsa repellendus fugit nisi",
        "completed": True
    },
    {
        "userId": 1,
        "id": 13,
        "title": "et doloremque nulla",
        "completed": False
    },
    {
        "userId": 1,
        "id": 14,
        "title": "repellendus sunt dolores architecto voluptatum",
        "completed": True
    },
    {
        "userId": 1,
        "id": 15,
        "title": "ab voluptatum amet voluptas",
        "completed": True
    },
    {
        "userId": 1,
        "id": 16,
        "title": "accusamus eos facilis sint et aut voluptatem",
        "completed": True
    },
    {
        "userId": 1,
        "id": 17,
        "title": "quo laboriosam deleniti aut qui",
        "completed": True
    },
    {
        "userId": 1,
        "id": 18,
        "title": "dolorum est consequatur ea mollitia in culpa",
        "completed": False
    },
    {
        "userId": 1,
        "id": 19,
        "title": "molestiae ipsa aut voluptatibus pariatur dolor nihil",
        "completed": True
    },
    {
        "userId": 1,
        "id": 20,
        "title": "ullam nobis libero sapiente ad optio sint",
        "completed": True
    },
    {
        "userId": 2,
        "id": 21,
        "title": "suscipit repellat esse quibusdam voluptatem incidunt",
        "completed": False
    },
    {
        "userId": 2,
        "id": 22,
        "title": "distinctio vitae autem nihil ut molestias quo",
        "completed": True
    },
    {
        "userId": 2,
        "id": 23,
        "title": "et itaque necessitatibus maxime molestiae qui quas velit",
        "completed": False
    },
    {
        "userId": 2,
        "id": 24,
        "title": "adipisci non ad dicta qui amet quaerat doloribus ea",
        "completed": False
    },
    {
        "userId": 2,
        "id": 25,
        "title": "voluptas quo tenetur perspiciatis explicabo natus",
        "completed": True
    },
    {
        "userId": 2,
        "id": 26,
        "title": "aliquam aut quasi",
        "completed": True
    },
    {
        "userId": 2,
        "id": 27,
        "title": "veritatis pariatur delectus",
        "completed": True
    },
    {
        "userId": 2,
        "id": 28,
        "title": "nesciunt totam sit blanditiis sit",
        "completed": False
    },
    {
        "userId": 2,
        "id": 29,
        "title": "laborum aut in quam",
        "completed": False
    },
    {
        "userId": 2,
        "id": 30,
        "title": "nemo perspiciatis repellat ut dolor libero commodi blanditiis omnis",
        "completed": True
    },
    {
        "userId": 2,
        "id": 31,
        "title": "repudiandae totam in est sint facere fuga",
        "completed": False
    },
    {
        "userId": 2,
        "id": 32,
        "title": "earum doloribus ea doloremque quis",
        "completed": False
    },
    {
        "userId": 2,
        "id": 33,
        "title": "sint sit aut vero",
        "completed": False
    },
    {
        "userId": 2,
        "id": 34,
        "title": "porro aut necessitatibus eaque distinctio",
        "completed": False
    },
    {
        "userId": 2,
        "id": 35,
        "title": "repellendus veritatis molestias dicta incidunt",
        "completed": True
    },
    {
        "userId": 2,
        "id": 36,
        "title": "excepturi deleniti adipisci voluptatem et neque optio illum ad",
        "completed": True
    },
    {
        "userId": 2,
        "id": 37,
        "title": "sunt cum tempora",
        "completed": False
    },
    {
        "userId": 2,
        "id": 38,
        "title": "totam quia non",
        "completed": False
    },
    {
        "userId": 2,
        "id": 39,
        "title": "doloremque quibusdam asperiores libero corrupti illum qui omnis",
        "completed": False
    },
    {
        "userId": 2,
        "id": 40,
        "title": "totam atque quo nesciunt",
        "completed": True
    },
    {
        "userId": 3,
        "id": 41,
        "title": "aliquid amet impedit consequatur aspernatur placeat eaque fugiat suscipit",
        "completed": False
    },
    {
        "userId": 3,
        "id": 42,
        "title": "rerum perferendis error quia ut eveniet",
        "completed": False
    },
    {
        "userId": 3,
        "id": 43,
        "title": "tempore ut sint quis recusandae",
        "completed": True
    },
    {
        "userId": 3,
        "id": 44,
        "title": "cum debitis quis accusamus doloremque ipsa natus sapiente omnis",
        "completed": True
    },
    {
        "userId": 3,
        "id": 45,
        "title": "velit soluta adipisci molestias reiciendis harum",
        "completed": False
    },
    {
        "userId": 3,
        "id": 46,
        "title": "vel voluptatem repellat nihil placeat corporis",
        "completed": False
    },
    {
        "userId": 3,
        "id": 47,
        "title": "nam qui rerum fugiat accusamus",
        "completed": False
    },
    {
        "userId": 3,
        "id": 48,
        "title": "sit reprehenderit omnis quia",
        "completed": False
    },
    {
        "userId": 3,
        "id": 49,
        "title": "ut necessitatibus aut maiores debitis officia blanditiis velit et",
        "completed": False
    },
    {
        "userId": 3,
        "id": 50,
        "title": "cupiditate necessitatibus ullam aut quis dolor voluptate",
        "completed": True
    },
    {
        "userId": 3,
        "id": 51,
        "title": "distinctio exercitationem ab doloribus",
        "completed": False
    },
    {
        "userId": 3,
        "id": 52,
        "title": "nesciunt dolorum quis recusandae ad pariatur ratione",
        "completed": False
    },
    {
        "userId": 3,
        "id": 53,
        "title": "qui labore est occaecati recusandae aliquid quam",
        "completed": False
    },
    {
        "userId": 3,
        "id": 54,
        "title": "quis et est ut voluptate quam dolor",
        "completed": True
    },
    {
        "userId": 3,
        "id": 55,
        "title": "voluptatum omnis minima qui occaecati provident nulla voluptatem ratione",
        "completed": True
    },
    {
        "userId": 3,
        "id": 56,
        "title": "deleniti ea temporibus enim",
        "completed": True
    },
    {
        "userId": 3,
        "id": 57,
        "title": "pariatur et magnam ea doloribus similique voluptatem rerum quia",
        "completed": False
    },
    {
        "userId": 3,
        "id": 58,
        "title": "est dicta totam qui explicabo doloribus qui dignissimos",
        "completed": False
    },
    {
        "userId": 3,
        "id": 59,
        "title": "perspiciatis velit id laborum placeat iusto et aliquam odio",
        "completed": False
    },
    {
        "userId": 3,
        "id": 60,
        "title": "et sequi qui architecto ut adipisci",
        "completed": True
    },
    {
        "userId": 4,
        "id": 61,
        "title": "odit optio omnis qui sunt",
        "completed": True
    },
    {
        "userId": 4,
        "id": 62,
        "title": "et placeat et tempore aspernatur sint numquam",
        "completed": False
    },
    {
        "userId": 4,
        "id": 63,
        "title": "doloremque aut dolores quidem fuga qui nulla",
        "completed": True
    },
    {
        "userId": 4,
        "id": 64,
        "title": "voluptas consequatur qui ut quia magnam nemo esse",
        "completed": False
    },
    {
        "userId": 4,
        "id": 65,
        "title": "fugiat pariatur ratione ut asperiores necessitatibus magni",
        "completed": False
    },
    {
        "userId": 4,
        "id": 66,
        "title": "rerum eum molestias autem voluptatum sit optio",
        "completed": False
    },
    {
        "userId": 4,
        "id": 67,
        "title": "quia voluptatibus voluptatem quos similique maiores repellat",
        "completed": False
    },
    {
        "userId": 4,
        "id": 68,
        "title": "aut id perspiciatis voluptatem iusto",
        "completed": False
    },
    {
        "userId": 4,
        "id": 69,
        "title": "doloribus sint dolorum ab adipisci itaque dignissimos aliquam suscipit",
        "completed": False
    },
    {
        "userId": 4,
        "id": 70,
        "title": "ut sequi accusantium et mollitia delectus sunt",
        "completed": False
    },
    {
        "userId": 4,
        "id": 71,
        "title": "aut velit saepe ullam",
        "completed": False
    },
    {
        "userId": 4,
        "id": 72,
        "title": "praesentium facilis facere quis harum voluptatibus voluptatem eum",
        "completed": False
    },
    {
        "userId": 4,
        "id": 73,
        "title": "sint amet quia totam corporis qui exercitationem commodi",
        "completed": True
    },
    {
        "userId": 4,
        "id": 74,
        "title": "expedita tempore nobis eveniet laborum maiores",
        "completed": False
    },
    {
        "userId": 4,
        "id": 75,
        "title": "occaecati adipisci est possimus totam",
        "completed": False
    },
    {
        "userId": 4,
        "id": 76,
        "title": "sequi dolorem sed",
        "completed": True
    },
    {
        "userId": 4,
        "id": 77,
        "title": "maiores aut nesciunt delectus exercitationem vel assumenda eligendi at",
        "completed": False
    },
    {
        "userId": 4,
        "id": 78,
        "title": "reiciendis est magnam amet nemo iste recusandae impedit quaerat",
        "completed": False
    },
    {
        "userId": 4,
        "id": 79,
        "title": "eum ipsa maxime ut",
        "completed": True
    },
    {
        "userId": 4,
        "id": 80,
        "title": "tempore molestias dolores rerum sequi voluptates ipsum consequatur",
        "completed": True
    },
    {
        "userId": 5,
        "id": 81,
        "title": "suscipit qui totam",
        "completed": True
    },
    {
        "userId": 5,
        "id": 82,
        "title": "voluptates eum voluptas et dicta",
        "completed": False
    },
    {
        "userId": 5,
        "id": 83,
        "title": "quidem at rerum quis ex aut sit quam",
        "completed": True
    },
    {
        "userId": 5,
        "id": 84,
        "title": "sunt veritatis ut voluptate",
        "completed": False
    },
    {
        "userId": 5,
        "id": 85,
        "title": "et quia ad iste a",
        "completed": True
    },
    {
        "userId": 5,
        "id": 86,
        "title": "incidunt ut saepe autem",
        "completed": True
    },
    {
        "userId": 5,
        "id": 87,
        "title": "laudantium quae eligendi consequatur quia et vero autem",
        "completed": True
    },
    {
        "userId": 5,
        "id": 88,
        "title": "vitae aut excepturi laboriosam sint aliquam et et accusantium",
        "completed": False
    },
    {
        "userId": 5,
        "id": 89,
        "title": "sequi ut omnis et",
        "completed": True
    },
    {
        "userId": 5,
        "id": 90,
        "title": "molestiae nisi accusantium tenetur dolorem et",
        "completed": True
    },
    {
        "userId": 5,
        "id": 91,
        "title": "nulla quis consequatur saepe qui id expedita",
        "completed": True
    },
    {
        "userId": 5,
        "id": 92,
        "title": "in omnis laboriosam",
        "completed": True
    },
    {
        "userId": 5,
        "id": 93,
        "title": "odio iure consequatur molestiae quibusdam necessitatibus quia sint",
        "completed": True
    },
    {
        "userId": 5,
        "id": 94,
        "title": "facilis modi saepe mollitia",
        "completed": False
    },
    {
        "userId": 5,
        "id": 95,
        "title": "vel nihil et molestiae iusto assumenda nemo quo ut",
        "completed": True
    },
    {
        "userId": 5,
        "id": 96,
        "title": "nobis suscipit ducimus enim asperiores voluptas",
        "completed": False
    },
    {
        "userId": 5,
        "id": 97,
        "title": "dolorum laboriosam eos qui iure aliquam",
        "completed": False
    },
    {
        "userId": 5,
        "id": 98,
        "title": "debitis accusantium ut quo facilis nihil quis sapiente necessitatibus",
        "completed": True
    },
    {
        "userId": 5,
        "id": 99,
        "title": "neque voluptates ratione",
        "completed": False
    },
    {
        "userId": 5,
        "id": 100,
        "title": "excepturi a et neque qui expedita vel voluptate",
        "completed": False
    },
    {
        "userId": 6,
        "id": 101,
        "title": "explicabo enim cumque porro aperiam occaecati minima",
        "completed": False
    },
    {
        "userId": 6,
        "id": 102,
        "title": "sed ab consequatur",
        "completed": False
    },
    {
        "userId": 6,
        "id": 103,
        "title": "non sunt delectus illo nulla tenetur enim omnis",
        "completed": False
    },
    {
        "userId": 6,
        "id": 104,
        "title": "excepturi non laudantium quo",
        "completed": False
    },
    {
        "userId": 6,
        "id": 105,
        "title": "totam quia dolorem et illum repellat voluptas optio",
        "completed": True
    },
    {
        "userId": 6,
        "id": 106,
        "title": "ad illo quis voluptatem temporibus",
        "completed": True
    },
    {
        "userId": 6,
        "id": 107,
        "title": "praesentium facilis omnis laudantium fugit ad iusto nihil nesciunt",
        "completed": False
    },
    {
        "userId": 6,
        "id": 108,
        "title": "a eos eaque nihil et exercitationem incidunt delectus",
        "completed": True
    },
    {
        "userId": 6,
        "id": 109,
        "title": "autem temporibus harum quisquam in culpa",
        "completed": True
    },
    {
        "userId": 6,
        "id": 110,
        "title": "aut aut ea corporis",
        "completed": True
    },
    {
        "userId": 6,
        "id": 111,
        "title": "magni accusantium labore et id quis provident",
        "completed": False
    },
    {
        "userId": 6,
        "id": 112,
        "title": "consectetur impedit quisquam qui deserunt non rerum consequuntur eius",
        "completed": False
    },
    {
        "userId": 6,
        "id": 113,
        "title": "quia atque aliquam sunt impedit voluptatum rerum assumenda nisi",
        "completed": False
    },
    {
        "userId": 6,
        "id": 114,
        "title": "cupiditate quos possimus corporis quisquam exercitationem beatae",
        "completed": False
    },
    {
        "userId": 6,
        "id": 115,
        "title": "sed et ea eum",
        "completed": False
    },
    {
        "userId": 6,
        "id": 116,
        "title": "ipsa dolores vel facilis ut",
        "completed": True
    },
    {
        "userId": 6,
        "id": 117,
        "title": "sequi quae est et qui qui eveniet asperiores",
        "completed": False
    },
    {
        "userId": 6,
        "id": 118,
        "title": "quia modi consequatur vero fugiat",
        "completed": False
    },
    {
        "userId": 6,
        "id": 119,
        "title": "corporis ducimus ea perspiciatis iste",
        "completed": False
    },
    {
        "userId": 6,
        "id": 120,
        "title": "dolorem laboriosam vel voluptas et aliquam quasi",
        "completed": False
    },
    {
        "userId": 7,
        "id": 121,
        "title": "inventore aut nihil minima laudantium hic qui omnis",
        "completed": True
    },
    {
        "userId": 7,
        "id": 122,
        "title": "provident aut nobis culpa",
        "completed": True
    },
    {
        "userId": 7,
        "id": 123,
        "title": "esse et quis iste est earum aut impedit",
        "completed": False
    },
    {
        "userId": 7,
        "id": 124,
        "title": "qui consectetur id",
        "completed": False
    },
    {
        "userId": 7,
        "id": 125,
        "title": "aut quasi autem iste tempore illum possimus",
        "completed": False
    },
    {
        "userId": 7,
        "id": 126,
        "title": "ut asperiores perspiciatis veniam ipsum rerum saepe",
        "completed": True
    },
    {
        "userId": 7,
        "id": 127,
        "title": "voluptatem libero consectetur rerum ut",
        "completed": True
    },
    {
        "userId": 7,
        "id": 128,
        "title": "eius omnis est qui voluptatem autem",
        "completed": False
    },
    {
        "userId": 7,
        "id": 129,
        "title": "rerum culpa quis harum",
        "completed": False
    },
    {
        "userId": 7,
        "id": 130,
        "title": "nulla aliquid eveniet harum laborum libero alias ut unde",
        "completed": True
    },
    {
        "userId": 7,
        "id": 131,
        "title": "qui ea incidunt quis",
        "completed": False
    },
    {
        "userId": 7,
        "id": 132,
        "title": "qui molestiae voluptatibus velit iure harum quisquam",
        "completed": True
    },
    {
        "userId": 7,
        "id": 133,
        "title": "et labore eos enim rerum consequatur sunt",
        "completed": True
    },
    {
        "userId": 7,
        "id": 134,
        "title": "molestiae doloribus et laborum quod ea",
        "completed": False
    },
    {
        "userId": 7,
        "id": 135,
        "title": "facere ipsa nam eum voluptates reiciendis vero qui",
        "completed": False
    },
    {
        "userId": 7,
        "id": 136,
        "title": "asperiores illo tempora fuga sed ut quasi adipisci",
        "completed": False
    },
    {
        "userId": 7,
        "id": 137,
        "title": "qui sit non",
        "completed": False
    },
    {
        "userId": 7,
        "id": 138,
        "title": "placeat minima consequatur rem qui ut",
        "completed": True
    },
    {
        "userId": 7,
        "id": 139,
        "title": "consequatur doloribus id possimus voluptas a voluptatem",
        "completed": False
    },
    {
        "userId": 7,
        "id": 140,
        "title": "aut consectetur in blanditiis deserunt quia sed laboriosam",
        "completed": True
    },
    {
        "userId": 8,
        "id": 141,
        "title": "explicabo consectetur debitis voluptates quas quae culpa rerum non",
        "completed": True
    },
    {
        "userId": 8,
        "id": 142,
        "title": "maiores accusantium architecto necessitatibus reiciendis ea aut",
        "completed": True
    },
    {
        "userId": 8,
        "id": 143,
        "title": "eum non recusandae cupiditate animi",
        "completed": False
    },
    {
        "userId": 8,
        "id": 144,
        "title": "ut eum exercitationem sint",
        "completed": False
    },
    {
        "userId": 8,
        "id": 145,
        "title": "beatae qui ullam incidunt voluptatem non nisi aliquam",
        "completed": False
    },
    {
        "userId": 8,
        "id": 146,
        "title": "molestiae suscipit ratione nihil odio libero impedit vero totam",
        "completed": True
    },
    {
        "userId": 8,
        "id": 147,
        "title": "eum itaque quod reprehenderit et facilis dolor autem ut",
        "completed": True
    },
    {
        "userId": 8,
        "id": 148,
        "title": "esse quas et quo quasi exercitationem",
        "completed": False
    },
    {
        "userId": 8,
        "id": 149,
        "title": "animi voluptas quod perferendis est",
        "completed": False
    },
    {
        "userId": 8,
        "id": 150,
        "title": "eos amet tempore laudantium fugit a",
        "completed": False
    },
    {
        "userId": 8,
        "id": 151,
        "title": "accusamus adipisci dicta qui quo ea explicabo sed vero",
        "completed": True
    },
    {
        "userId": 8,
        "id": 152,
        "title": "odit eligendi recusandae doloremque cumque non",
        "completed": False
    },
    {
        "userId": 8,
        "id": 153,
        "title": "ea aperiam consequatur qui repellat eos",
        "completed": False
    },
    {
        "userId": 8,
        "id": 154,
        "title": "rerum non ex sapiente",
        "completed": True
    },
    {
        "userId": 8,
        "id": 155,
        "title": "voluptatem nobis consequatur et assumenda magnam",
        "completed": True
    },
    {
        "userId": 8,
        "id": 156,
        "title": "nam quia quia nulla repellat assumenda quibusdam sit nobis",
        "completed": True
    },
    {
        "userId": 8,
        "id": 157,
        "title": "dolorem veniam quisquam deserunt repellendus",
        "completed": True
    },
    {
        "userId": 8,
        "id": 158,
        "title": "debitis vitae delectus et harum accusamus aut deleniti a",
        "completed": True
    },
    {
        "userId": 8,
        "id": 159,
        "title": "debitis adipisci quibusdam aliquam sed dolore ea praesentium nobis",
        "completed": True
    },
    {
        "userId": 8,
        "id": 160,
        "title": "et praesentium aliquam est",
        "completed": False
    },
    {
        "userId": 9,
        "id": 161,
        "title": "ex hic consequuntur earum omnis alias ut occaecati culpa",
        "completed": True
    },
    {
        "userId": 9,
        "id": 162,
        "title": "omnis laboriosam molestias animi sunt dolore",
        "completed": True
    },
    {
        "userId": 9,
        "id": 163,
        "title": "natus corrupti maxime laudantium et voluptatem laboriosam odit",
        "completed": False
    },
    {
        "userId": 9,
        "id": 164,
        "title": "reprehenderit quos aut aut consequatur est sed",
        "completed": False
    },
    {
        "userId": 9,
        "id": 165,
        "title": "fugiat perferendis sed aut quidem",
        "completed": False
    },
    {
        "userId": 9,
        "id": 166,
        "title": "quos quo possimus suscipit minima ut",
        "completed": False
    },
    {
        "userId": 9,
        "id": 167,
        "title": "et quis minus quo a asperiores molestiae",
        "completed": False
    },
    {
        "userId": 9,
        "id": 168,
        "title": "recusandae quia qui sunt libero",
        "completed": False
    },
    {
        "userId": 9,
        "id": 169,
        "title": "ea odio perferendis officiis",
        "completed": True
    },
    {
        "userId": 9,
        "id": 170,
        "title": "quisquam aliquam quia doloribus aut",
        "completed": False
    },
    {
        "userId": 9,
        "id": 171,
        "title": "fugiat aut voluptatibus corrupti deleniti velit iste odio",
        "completed": True
    },
    {
        "userId": 9,
        "id": 172,
        "title": "et provident amet rerum consectetur et voluptatum",
        "completed": False
    },
    {
        "userId": 9,
        "id": 173,
        "title": "harum ad aperiam quis",
        "completed": False
    },
    {
        "userId": 9,
        "id": 174,
        "title": "similique aut quo",
        "completed": False
    },
    {
        "userId": 9,
        "id": 175,
        "title": "laudantium eius officia perferendis provident perspiciatis asperiores",
        "completed": True
    },
    {
        "userId": 9,
        "id": 176,
        "title": "magni soluta corrupti ut maiores rem quidem",
        "completed": False
    },
    {
        "userId": 9,
        "id": 177,
        "title": "et placeat temporibus voluptas est tempora quos quibusdam",
        "completed": False
    },
    {
        "userId": 9,
        "id": 178,
        "title": "nesciunt itaque commodi tempore",
        "completed": True
    },
    {
        "userId": 9,
        "id": 179,
        "title": "omnis consequuntur cupiditate impedit itaque ipsam quo",
        "completed": True
    },
    {
        "userId": 9,
        "id": 180,
        "title": "debitis nisi et dolorem repellat et",
        "completed": True
    },
    {
        "userId": 10,
        "id": 181,
        "title": "ut cupiditate sequi aliquam fuga maiores",
        "completed": False
    },
    {
        "userId": 10,
        "id": 182,
        "title": "inventore saepe cumque et aut illum enim",
        "completed": True
    },
    {
        "userId": 10,
        "id": 183,
        "title": "omnis nulla eum aliquam distinctio",
        "completed": True
    },
    {
        "userId": 10,
        "id": 184,
        "title": "molestias modi perferendis perspiciatis",
        "completed": False
    },
    {
        "userId": 10,
        "id": 185,
        "title": "voluptates dignissimos sed doloribus animi quaerat aut",
        "completed": False
    },
    {
        "userId": 10,
        "id": 186,
        "title": "explicabo odio est et",
        "completed": False
    },
    {
        "userId": 10,
        "id": 187,
        "title": "consequuntur animi possimus",
        "completed": False
    },
    {
        "userId": 10,
        "id": 188,
        "title": "vel non beatae est",
        "completed": True
    },
    {
        "userId": 10,
        "id": 189,
        "title": "culpa eius et voluptatem et",
        "completed": True
    },
    {
        "userId": 10,
        "id": 190,
        "title": "accusamus sint iusto et voluptatem exercitationem",
        "completed": True
    },
    {
        "userId": 10,
        "id": 191,
        "title": "temporibus atque distinctio omnis eius impedit tempore molestias pariatur",
        "completed": True
    },
    {
        "userId": 10,
        "id": 192,
        "title": "ut quas possimus exercitationem sint voluptates",
        "completed": False
    },
    {
        "userId": 10,
        "id": 193,
        "title": "rerum debitis voluptatem qui eveniet tempora distinctio a",
        "completed": True
    },
    {
        "userId": 10,
        "id": 194,
        "title": "sed ut vero sit molestiae",
        "completed": False
    },
    {
        "userId": 10,
        "id": 195,
        "title": "rerum ex veniam mollitia voluptatibus pariatur",
        "completed": True
    },
    {
        "userId": 10,
        "id": 196,
        "title": "consequuntur aut ut fugit similique",
        "completed": True
    },
    {
        "userId": 10,
        "id": 197,
        "title": "dignissimos quo nobis earum saepe",
        "completed": True
    },
    {
        "userId": 10,
        "id": 198,
        "title": "quis eius est sint explicabo",
        "completed": True
    },
    {
        "userId": 10,
        "id": 199,
        "title": "numquam repellendus a magnam",
        "completed": True
    },
    {
        "userId": 10,
        "id": 200,
        "title": "ipsam aperiam voluptates qui",
        "completed": False
    }
]

with open('date.json', 'w') as f:
    json.dump(date, f)

with open('date.json', 'r') as f:
    
    date = json.load(f)


for i, d in enumerate(date,date+1):
    
    with open(f'dict{i}.json', 'w') as f:
        json.dump(d, f)



from docx import Document

doc = Document()
doc.add_paragraph('Hello Python')
doc.save('file1.docx')

doc = Document('file1.docx')
for para in doc.paragraphs:
    for run in para.runs:
        if run.bold:
            print(run.text)

doc = Document()
run = doc.add_paragraph().add_run('Hello Python')
run.bold = True
doc.save('file2.docx')